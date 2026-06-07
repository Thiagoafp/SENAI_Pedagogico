"""
gerador_fo.py
Gerador dedicado para o Plano de Ensino SENAI (FO-178).
"""

import io
import copy
import json
import re
import sqlite3
import anthropic
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from extrator_ementa import carregar_uc, DB_PATH

MODEL = "claude-sonnet-4-5"


# ─── Helpers de formatação ────────────────────────────────────────────────────

def set_cell_text(cell, text: str, bold: bool = False, italic: bool = False,
                  font_size: int = 10, uppercase: bool = False,
                  align: str = "left", valign: str = "center"):
    """Define texto da célula com Arial, formatação e alinhamento."""
    if uppercase:
        text = text.upper()

    # Limpa parágrafos extras
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    para = cell.paragraphs[0]

    # Limpa runs existentes
    for run in para.runs:
        run.text = ""

    # Alinhamento parágrafo
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    # Cria run
    run = para.runs[0] if para.runs else para.add_run(text)
    run.text = text
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = "Arial"

    # Força Arial no XML
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

    # Alinhamento vertical
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), valign)
    existing_va = tcPr.find(qn("w:vAlign"))
    if existing_va is not None:
        tcPr.remove(existing_va)
    tcPr.append(vAlign)


def set_cell_empty(cell):
    """Limpa célula completamente sem texto."""
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if not para.runs:
        para.add_run("")


def set_vmerge(cell, tipo: str):
    """
    Define vMerge na célula.
    tipo: 'restart' = início do merge, 'continue' = continuação (célula vazia)
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:vMerge"))
    if existing is not None:
        tcPr.remove(existing)

    vMerge = OxmlElement("w:vMerge")
    if tipo == "restart":
        vMerge.set(qn("w:val"), "restart")
    # continue = sem atributo val
    tcPr.append(vMerge)

    # Célula continue deve estar vazia
    if tipo == "continue":
        set_cell_empty(cell)


def set_valign_table(table, valign: str = "center"):
    """Alinhamento vertical para todas as células — via XML direto."""
    for row in table.rows:
        # Acessa _tc diretamente sem resolver merge
        for tc in row._tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            existing = tcPr.find(qn("w:vAlign"))
            if existing is not None:
                tcPr.remove(existing)
            vAlign = OxmlElement("w:vAlign")
            vAlign.set(qn("w:val"), valign)
            tcPr.append(vAlign)


def calcular_bimestre(num_aula: int, total_aulas: int) -> str:
    por_bim = max(total_aulas / 4, 1)
    bim = min(int((num_aula - 1) / por_bim) + 1, 4)
    return f"{bim}º BIM"


# ─── Geração de conteúdo — duas chamadas separadas ───────────────────────────

def gerar_cabecalho_fo(uc: dict, contexto: str, api_key: str) -> dict:
    """
    Chamada 1: gera dados do cabeçalho (estratégia, contextualização, desafio, resultados).
    Textos mais curtos e focados.
    """
    client = anthropic.Anthropic(api_key=api_key)

    capacidades = "\n".join([f"- {c}" for c in uc["capacidades"]])
    conhecimentos = "\n".join([f"- {c}" for c in uc["conhecimentos"]])

    # Verifica campos vazios no banco — IA gera quando necessario
    refs_banco = uc.get("referencias", [])
    objetivo_banco = uc.get("objetivo_geral", "").strip()

    instrucao_extras = ""
    if not objetivo_banco:
        instrucao_extras += """,
  "objetivo_geral": "Gere o objetivo geral desta UC em uma frase clara descrevendo a competencia que o aluno desenvolve (maximo 300 caracteres)" """
    if not refs_banco:
        instrucao_extras += """,
  "referencias": "Gere 3 referencias bibliograficas relevantes para esta UC no formato ABNT. Use autores e obras reais conhecidas do mercado e da academia para este tema." """

    prompt = f"""Voce e especialista pedagogico do SENAI preenchendo o cabecalho de um Plano de Ensino.

UC: {uc['nome']}
Objetivo Geral no banco: {objetivo_banco or 'NAO INFORMADO - voce deve gerar no campo objetivo_geral'}
Capacidades: {capacidades}
Contexto adicional do professor: {contexto or 'Sem contexto adicional.'}

Retorne APENAS JSON valido sem markdown:

{{
  "estrategia_principal": "Projeto",
  "contextualizacao": "2 paragrafos curtos contextualizando esta UC com a realidade dos alunos (MAXIMO 400 caracteres no total, seja breve)",
  "desafio": "1 paragrafo descrevendo o desafio central que os alunos resolverao (MAXIMO 300 caracteres, seja breve)",
  "resultados_esperados": "Topicos do que os alunos serao capazes ao final (MAXIMO 300 caracteres, seja breve)"{instrucao_extras}
}}

estrategia_principal deve ser exatamente uma das opcoes: "Estudo de caso", "Projeto", "Situacao-Problema" ou "Pesquisa Aplicada"
CRITICO: textos CURTOS, respeite os limites de caracteres para nao ultrapassar o espaco da tabela no documento."""

    msg = client.messages.create(
        model=MODEL, max_tokens=2000,
        messages=[{"role": "user", "content": prompt}]
    )

    texto = msg.content[0].text.strip()
    texto = re.sub(r'^```json\s*', '', texto)
    texto = re.sub(r'\s*```$', '', texto)
    return json.loads(texto)


def gerar_aulas_fo(uc: dict, contexto: str, api_key: str, total_encontros: int = None) -> list:
    """
    Chamada 2: gera apenas a lista de encontros/aulas.
    Dividida em blocos de 20 para evitar JSON invalido.
    total_encontros: numero de linhas do FO (pode ser menor que total H/A para UCs com multiplas h/encontro)
    """
    client = anthropic.Anthropic(api_key=api_key)
    total_ha = int(str(uc.get("carga_horaria_ha", "40")).replace("H/A", "").strip() or 40)
    if total_encontros is None:
        total_encontros = total_ha
    horas_por_encontro = total_ha // total_encontros
    capacidades = "\n".join([f"- {c}" for c in uc["capacidades"]])
    conhecimentos = "\n".join([f"- {c}" for c in uc["conhecimentos"]])

    todas_aulas = []

    # Divide em blocos de 20
    blocos = [(1, min(20, total_encontros))]
    if total_encontros > 20:
        blocos.append((21, total_encontros))

    for inicio, fim in blocos:
        prompt = f"""Gere APENAS a lista de encontros {inicio} a {fim} para o Plano de Ensino SENAI.

UC: {uc['nome']} | Total: {total_ha} H/A em {total_encontros} encontros de {horas_por_encontro}h cada
Capacidades: {capacidades}
Conhecimentos disponiveis: {conhecimentos}
Contexto: {contexto or 'Sem contexto.'}

Retorne APENAS JSON valido sem markdown:

{{
  "aulas": [
    {{
      "numero": {inicio},
      "capacidade": "capacidade desenvolvida neste encontro",
      "conhecimento": "Titulo do conteudo abordado",
      "estrategia": "Teoria (Projetor/Lousa); Exercicios Guiados; Laboratorio.",
      "criterio_avaliacao": "como sera avaliado",
      "instrumento": "instrumento especifico",
      "recursos": "recursos utilizados"
    }}
  ]
}}

Gere exatamente os encontros de {inicio} ate {fim}.
Textos curtos e objetivos em cada campo (maximo 100 caracteres por campo).
Progressao logica dos conteudos cobrindo toda a UC em {total_encontros} encontros."""

        msg = client.messages.create(
            model=MODEL, max_tokens=8000,
            messages=[{"role": "user", "content": prompt}]
        )

        texto = msg.content[0].text.strip()
        texto = re.sub(r'^```json\s*', '', texto)
        texto = re.sub(r'\s*```$', '', texto)

        try:
            dados = json.loads(texto)
            todas_aulas.extend(dados["aulas"])
        except json.JSONDecodeError:
            # Tenta extrair JSON do texto
            match = re.search(r'\{.*\}', texto, re.DOTALL)
            if match:
                dados = json.loads(match.group())
                todas_aulas.extend(dados["aulas"])
            else:
                raise ValueError(f"JSON inválido no bloco {inicio}-{fim}. Tente novamente.")

    return todas_aulas


# ─── Preenchedor do template ──────────────────────────────────────────────────

def preencher_fo(template_path: str, uc: dict, cabecalho: dict,
                 aulas: list, nome_docente: str = "",
                 funcao: str = "Instrutor de Informatica",
                 subfuncao: str = "") -> bytes:
    doc = Document(template_path)
    tables = doc.tables
    total_aulas = len(aulas)
    total_ha = int(str(uc.get("carga_horaria_ha", "40")).replace("H/A", "").strip() or 40)
    horas_por_encontro = max(1, total_ha // total_aulas) if total_aulas > 0 else 1

    # ── TABELA 0: Cabeçalho ───────────────────────────────────────────────────
    t0 = tables[0]
    set_valign_table(t0, "center")
    set_cell_text(t0.rows[0].cells[1],
                  uc.get('curso_nome', 'TÉCNICO EM DESENVOLVIMENTO DE SISTEMAS'),
                  bold=True, uppercase=True, font_size=10)
    set_cell_text(t0.rows[1].cells[1],
                  uc['nome'], bold=True, uppercase=True, font_size=10)
    set_cell_text(t0.rows[2].cells[1], nome_docente.upper() if nome_docente else '', font_size=10)
    set_cell_text(t0.rows[2].cells[3],
                  f"{uc.get('carga_horaria_ha', '40')}h", font_size=10)

    # ── TABELA 1: Perfil — acessa via XML direto ────────────────────────────
    t1 = tables[1]

    def set_text_t1(row_idx, col_idx, text, uppercase=False, font_size=10):
        tcs = t1.rows[row_idx]._tr.findall(qn("w:tc"))
        if col_idx >= len(tcs):
            return
        tc = tcs[col_idx]
        if uppercase:
            text = text.upper()
        p = tc.find(qn("w:p"))
        if p is None:
            p = OxmlElement("w:p")
            tc.append(p)
        for r in p.findall(qn("w:r")):
            p.remove(r)
        r_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size * 2))
        rPr.append(rFonts)
        rPr.append(sz)
        r_el.append(rPr)
        t_el = OxmlElement("w:t")
        t_el.text = text
        if " " in text:
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_el.append(t_el)
        p.append(r_el)

    print(f"DEBUG funcao: '{funcao}'")
    print(f"DEBUG subfuncao: '{subfuncao}'")
    print(f"DEBUG objetivo: '{uc.get('objetivo_geral', '')[:50]}'")
    print(f"DEBUG refs: {uc.get('referencias', [])[:1]}")
    set_text_t1(1, 1, funcao, uppercase=True)
    set_text_t1(2, 1, subfuncao, uppercase=True)
    # Usa objetivo do banco ou o gerado pela IA
    objetivo = uc.get('objetivo_geral', '').strip() or cabecalho.get('objetivo_geral', '')
    set_text_t1(3, 1, objetivo, font_size=10)

    # ── TABELA 2: Estratégias ─────────────────────────────────────────────────
    t2 = tables[2]
    set_valign_table(t2, "center")
    opcoes = [
        ("Estudo de caso",    "(   ) Estudo de caso"),
        ("Projeto",           "(   ) Projeto (elaboração ou execução)"),
        ("Situação-Problema", "(   ) Situação-Problema"),
        ("Pesquisa Aplicada", "(   ) Pesquisa Aplicada"),
    ]
    principal = cabecalho.get("estrategia_principal", "Projeto")
    for i, (chave, label) in enumerate(opcoes):
        txt = label.replace("(   )", "( X )") if chave == principal else label
        set_cell_text(t2.rows[1].cells[i], txt,
                      font_size=10, align="center", valign="center")

    # ── TABELA 3: Contextualização / Desafio / Resultados ────────────────────
    # Estrutura correta de vMerge:
    #   [1,0] restart  [1,1] restart + conteúdo contextualização
    #   [2,0] continue [2,1] continue + vazio
    #   [3,0] restart  [3,1] restart + conteúdo desafio
    #   [4,0] continue [4,1] continue + vazio
    #   [5,0] none     [5,1] restart + conteúdo resultados
    #   [6,0] none     [6,1] continue + vazio

    t3 = tables[3]

    # Remove linha 6 se existir — Resultados Esperados não tem linha par
    while len(t3.rows) > 6:
        t3._tbl.remove(t3.rows[-1]._tr)

    # Alturas minimas da tabela 3 — atLeast permite crescer se necessario
    alturas_t3 = [280, 800, 800, 800, 800, 1200]
    for i, row in enumerate(t3.rows):
        if i >= len(alturas_t3):
            break
        trPr = row._tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            row._tr.insert(0, trPr)
        trH = trPr.find(qn("w:trHeight"))
        if trH is None:
            trH = OxmlElement("w:trHeight")
            trPr.append(trH)
        trH.set(qn("w:val"), str(alturas_t3[i]))
        trH.set(qn("w:hRule"), "atLeast")

    # Acessa células da tabela 3 diretamente pelo XML para evitar
    # erro de resolução de vMerge no python-docx

    def get_tc(table, row_idx, col_idx):
        """Retorna o _tc da célula sem resolver merge."""
        tcs = table.rows[row_idx]._tr.findall(qn("w:tc"))
        if col_idx < len(tcs):
            return tcs[col_idx]
        return None

    def set_vmerge_tc(tc, tipo):
        """Define vMerge diretamente no _tc."""
        if tc is None:
            return
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        existing = tcPr.find(qn("w:vMerge"))
        if existing is not None:
            tcPr.remove(existing)
        vMerge = OxmlElement("w:vMerge")
        if tipo == "restart":
            vMerge.set(qn("w:val"), "restart")
        tcPr.append(vMerge)
        if tipo == "continue":
            for p in tc.findall(qn("w:p")):
                for r in p.findall(qn("w:r")):
                    for t in r.findall(qn("w:t")):
                        t.text = ""

    def set_text_tc(tc, text, font_size=10, italic=False):
        """Define texto diretamente no _tc."""
        if tc is None:
            return
        # Pega ou cria parágrafo
        p = tc.find(qn("w:p"))
        if p is None:
            p = OxmlElement("w:p")
            tc.append(p)
        # Limpa runs existentes
        for r in p.findall(qn("w:r")):
            p.remove(r)
        # Cria novo run
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(font_size * 2))
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        rPr.append(rFonts)
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        if " " in text:
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)

    # Contextualização — linhas 1 e 2
    set_vmerge_tc(get_tc(t3, 1, 0), "restart")
    set_vmerge_tc(get_tc(t3, 1, 1), "restart")
    set_text_tc(get_tc(t3, 1, 1), cabecalho.get("contextualizacao", ""))

    set_vmerge_tc(get_tc(t3, 2, 0), "continue")
    set_vmerge_tc(get_tc(t3, 2, 1), "continue")

    # Desafio — linhas 3 e 4
    set_vmerge_tc(get_tc(t3, 3, 0), "restart")
    set_vmerge_tc(get_tc(t3, 3, 1), "restart")
    set_text_tc(get_tc(t3, 3, 1), cabecalho.get("desafio", ""))

    set_vmerge_tc(get_tc(t3, 4, 0), "continue")
    set_vmerge_tc(get_tc(t3, 4, 1), "continue")

    # Resultados — linha 5 apenas (sem par)
    set_vmerge_tc(get_tc(t3, 5, 1), "restart")
    set_text_tc(get_tc(t3, 5, 1), cabecalho.get("resultados_esperados", ""))

    # ── TABELA 4: Aulas ───────────────────────────────────────────────────────
    t4 = tables[4]
    linhas_necessarias = total_aulas + 1

    while len(t4.rows) < linhas_necessarias:
        ultima_tr = t4.rows[-1]._tr
        nova_tr = copy.deepcopy(ultima_tr)
        ultima_tr.addnext(nova_tr)

    while len(t4.rows) > linhas_necessarias:
        t4._tbl.remove(t4.rows[-1]._tr)

    for i, aula in enumerate(aulas):
        row = t4.rows[i + 1]
        bim = calcular_bimestre(aula["numero"], total_aulas)
        set_cell_text(row.cells[0], f"{aula['numero']}\n{bim}",
                      font_size=9, align="center", valign="center")
        horas_enc = total_ha // total_aulas if total_aulas > 0 else 1
        set_cell_text(row.cells[1], f"{horas_enc}h",
                      font_size=9, align="center", valign="center")
        set_cell_text(row.cells[2], aula.get("capacidade", ""),
                      font_size=9, valign="top")
        set_cell_text(row.cells[3], aula.get("conhecimento", ""),
                      font_size=9, valign="top")
        set_cell_text(row.cells[4], aula.get("estrategia", ""),
                      font_size=9, valign="top")
        set_cell_text(row.cells[5], aula.get("criterio_avaliacao", ""),
                      font_size=9, valign="top")
        set_cell_text(row.cells[6], aula.get("instrumento", ""),
                      font_size=9, valign="top")
        set_cell_text(row.cells[7], aula.get("recursos", ""),
                      font_size=9, valign="top")

    # ── TABELA 6: Referências — banco ou IA ─────────────────────────────────
    if len(tables) > 6:
        t6 = tables[6]
        refs_lista = uc.get("referencias", [])
        # Se banco tem referencias usa elas, senao usa as geradas pela IA
        if refs_lista:
            refs = "\n".join(refs_lista)
        else:
            refs = cabecalho.get("referencias", "")
        if len(t6.rows) > 1 and refs:
            tc_ref = t6.rows[1]._tr.findall(qn("w:tc"))
            if tc_ref:
                tc = tc_ref[0]
                p = tc.find(qn("w:p"))
                if p is None:
                    p = OxmlElement("w:p")
                    tc.append(p)
                for r in p.findall(qn("w:r")):
                    p.remove(r)
                r_el = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Arial")
                rFonts.set(qn("w:hAnsi"), "Arial")
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "18")
                rPr.append(rFonts)
                rPr.append(sz)
                r_el.append(rPr)
                t_el = OxmlElement("w:t")
                t_el.text = refs
                t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                r_el.append(t_el)
                p.append(r_el)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ─── Função principal ─────────────────────────────────────────────────────────

def gerar_fo_completo(uc_id: int, api_key: str, contexto: str = "",
                      nome_docente: str = "", template_path: str = None,
                      funcao: str = "Instrutor de Informatica",
                      subfuncao: str = "", on_etapa=None) -> bytes:
    """Gera o FO completo. on_etapa(descricao) chamado a cada etapa concluida."""
    uc = carregar_uc(uc_id)
    if not uc:
        raise ValueError(f"UC nao encontrada")

    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""SELECT nome FROM cursos WHERE id =
                   (SELECT curso_id FROM unidades_curriculares WHERE id = ?)""", (uc_id,))
    row = cur.fetchone()
    conn.close()
    uc['curso_nome'] = row[0] if row else "TECNICO EM DESENVOLVIMENTO DE SISTEMAS"

    if not template_path:
        template_path = "FO_template.docx"

    import streamlit as st

    # Cria ou reutiliza placeholder unico — sobrescreve sempre a mesma linha
    try:
        if "_fo_log_placeholder" not in st.session_state:
            st.session_state["_fo_log_placeholder"] = st.empty()
        _placeholder = st.session_state["_fo_log_placeholder"]
    except Exception:
        _placeholder = None

    def log(msg):
        try:
            msg_safe = msg.encode('ascii', 'ignore').decode('ascii').strip()
            if msg_safe and _placeholder:
                _placeholder.info(msg_safe)
        except Exception:
            print(msg)

    def notificar(descricao):
        if on_etapa:
            try:
                on_etapa(descricao)
            except Exception:
                pass

    log("Gerando cabecalho e estrategia...")
    cabecalho = gerar_cabecalho_fo(uc, contexto, api_key)
    log("Cabecalho gerado!")
    notificar("Cabecalho gerado")

    total_ha = int(str(uc.get("carga_horaria_ha", "40")).replace("H/A", "").strip() or 40)
    if total_ha <= 40:
        total_encontros = total_ha
    elif total_ha <= 80:
        total_encontros = total_ha // 2
    else:
        total_encontros = total_ha // 3
    log(f"Gerando {total_encontros} encontros de {total_ha // total_encontros}h...")
    aulas = gerar_aulas_fo(uc, contexto, api_key, total_encontros)
    log(f"{len(aulas)} encontros gerados!")
    notificar("Sequencia didatica gerada")

    log("Preenchendo template Word...")
    resultado = preencher_fo(template_path, uc, cabecalho, aulas,
                            nome_docente, funcao, subfuncao)
    log("FO finalizado!")
    notificar("Template preenchido")
    # Limpa o placeholder apos conclusao
    try:
        if _placeholder:
            _placeholder.empty()
        st.session_state.pop("_fo_log_placeholder", None)
    except Exception:
        pass
    return resultado


# ─── Limpeza de template ──────────────────────────────────────────────────────

def limpar_template_fo(conteudo_bytes: bytes) -> bytes:
    """Limpa dados preenchidos do template FO mantendo estrutura."""
    doc = Document(io.BytesIO(conteudo_bytes))

    # Tabela 0
    t0 = doc.tables[0]
    for run in t0.rows[0].cells[1].paragraphs[0].runs: run.text = ""
    for run in t0.rows[1].cells[1].paragraphs[0].runs: run.text = ""
    for run in t0.rows[2].cells[1].paragraphs[0].runs: run.text = ""
    for run in t0.rows[2].cells[3].paragraphs[0].runs: run.text = ""

    # Tabela 1
    t1 = doc.tables[1]
    for i in [1, 2, 3]:
        for run in t1.rows[i].cells[1].paragraphs[0].runs: run.text = ""

    # Tabela 2 — reset estratégias
    t2 = doc.tables[2]
    opcoes = ["(   ) Estudo de caso", "(   ) Projeto (elaboração ou execução)",
              "(   ) Situação-Problema", "(   ) Pesquisa Aplicada"]
    for i, opcao in enumerate(opcoes):
        for run in t2.rows[1].cells[i].paragraphs[0].runs: run.text = opcao

    # Tabela 3
    t3 = doc.tables[3]
    for row_idx in [1, 3, 5]:
        if row_idx < len(t3.rows):
            for para in t3.rows[row_idx].cells[1].paragraphs:
                for run in para.runs: run.text = ""

    # Tabela 4
    t4 = doc.tables[4]
    for row in t4.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.text = ""

    # Tabela 6
    t6 = doc.tables[6]
    for run in t6.rows[1].cells[0].paragraphs[0].runs: run.text = ""

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()