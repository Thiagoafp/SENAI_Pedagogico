"""
gerador_documentos.py
Usa Claude API + dados do SQLite para gerar os documentos pedagógicos.
Suporta templates personalizados (.docx) ou gera formato profissional padrão.
"""

import io
import re
import sqlite3
import anthropic
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from extrator_ementa import carregar_uc, DB_PATH

MODEL = "claude-sonnet-4-5"


# ─── Cliente Claude ───────────────────────────────────────────────────────────

def get_client(api_key: str) -> anthropic.Anthropic:
    return anthropic.Anthropic(api_key=api_key)


# ─── Helpers de template ──────────────────────────────────────────────────────

def extrair_texto_template(template_bytes: bytes) -> str:
    """Extrai o texto de um .docx template para usar no prompt."""
    try:
        doc = Document(io.BytesIO(template_bytes))
        linhas = []
        for para in doc.paragraphs:
            if para.text.strip():
                linhas.append(para.text)
        # Extrai tabelas também
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    linhas.append(row_text)
        return "\n".join(linhas)
    except Exception:
        return ""


def calcular_num_aulas(ch_ha: str, aulas_por_semana: int = 1) -> int:
    """
    Calcula número de aulas baseado na carga horária.
    ch_ha: carga horária em horas/aula (ex: "40")
    aulas_por_semana: quantas aulas por semana (padrão 1)
    Cada aula = 50min contado como 1 H/A
    """
    try:
        ch = int(str(ch_ha).replace("H/A", "").replace("h", "").strip())
        return ch  # 1 H/A = 1 aula de 50min
    except:
        return 20


# ─── Helpers de formatação do docx ───────────────────────────────────────────

def configurar_documento() -> Document:
    """Cria documento Word com configurações base."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    return doc


def adicionar_cabecalho(doc: Document, titulo: str, subtitulo: str = ""):
    """Adiciona cabeçalho padrão SENAI ao documento."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SENAI – Serviço Nacional de Aprendizagem Industrial")
    run.bold = True
    run.font.size = Pt(12)

    if subtitulo:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitulo)
        run2.font.size = Pt(11)

    doc.add_paragraph()
    h = doc.add_heading(titulo, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def doc_para_bytes(doc: Document) -> bytes:
    """Converte Document para bytes."""
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def texto_para_docx_padrao(conteudo: str, titulo: str, subtitulo: str) -> bytes:
    """Converte texto gerado pela IA em docx com formatação padrão."""
    doc = configurar_documento()
    adicionar_cabecalho(doc, titulo, subtitulo)

    linhas = conteudo.split("\n")
    for linha in linhas:
        linha_strip = linha.strip()
        if not linha_strip:
            doc.add_paragraph()
            continue

        # Detecta títulos de seção em maiúsculo ou com PARTE/CAPÍTULO
        if (linha_strip.isupper() and len(linha_strip) > 3) or \
           re.match(r'^(PARTE|CAPÍTULO|AULA|SEÇÃO|APRESENTAÇÃO|REFERÊNCIAS|GABARITO)', linha_strip.upper()):
            doc.add_heading(linha_strip, level=2)

        # Subtítulos numerados (1.1, 2.3, etc)
        elif re.match(r'^\d+\.\d+\s', linha_strip):
            doc.add_heading(linha_strip, level=3)

        # Campos com rótulo (Duração:, Objetivos:, etc)
        elif re.match(r'^[A-ZÀ-Ú][a-zà-ú\s]+:', linha_strip) and len(linha_strip.split(":")[0]) < 30:
            p = doc.add_paragraph()
            partes = linha_strip.split(":", 1)
            run = p.add_run(partes[0] + ": ")
            run.bold = True
            if len(partes) > 1 and partes[1].strip():
                p.add_run(partes[1].strip())

        # Alternativas de questão ( ) a)
        elif re.match(r'^\([xX ]\)\s*[a-d]\)', linha_strip) or re.match(r'^[a-d]\)', linha_strip):
            p = doc.add_paragraph(linha_strip)
            p.paragraph_format.left_indent = Cm(1)

        # Itens de lista com -
        elif linha_strip.startswith("- ") or linha_strip.startswith("• "):
            p = doc.add_paragraph(linha_strip[2:], style="List Bullet")

        # Separador ---
        elif linha_strip == "---":
            doc.add_paragraph("─" * 50)

        # Texto normal
        else:
            p = doc.add_paragraph(linha_strip)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc_para_bytes(doc)


# ─── Construtor de prompts ────────────────────────────────────────────────────

def montar_prompt(tipo: str, uc: dict, template_texto: str = "") -> str:
    """
    Monta o prompt para geração do documento.
    Se tiver template, inclui como referência de formato.
    """
    capacidades = "\n".join([f"- {c}" for c in uc["capacidades"]])
    conhecimentos = "\n".join([f"- {c}" for c in uc["conhecimentos"]])
    referencias = "\n".join([f"- {r}" for r in uc.get("referencias", [])])
    num_aulas = calcular_num_aulas(uc.get("carga_horaria_ha", "40"))

    # Instrução de formato
    if template_texto:
        instrucao_formato = f"""
O professor forneceu um template de documento que deve ser seguido.
Use EXATAMENTE o mesmo formato, estrutura e campos do template abaixo,
preenchendo com os dados reais da unidade curricular:

=== TEMPLATE DO PROFESSOR ===
{template_texto}
=== FIM DO TEMPLATE ===

Mantenha todos os campos, tabelas e estrutura do template.
Substitua apenas o conteúdo com as informações reais da UC.
"""
    else:
        instrucao_formato = "Crie um documento profissional e bem estruturado no padrão SENAI."

    # Dados da UC
    dados_uc = f"""
UNIDADE CURRICULAR: {uc['nome']}
Carga Horária: {uc.get('carga_horaria_ha', '')} H/A / {uc.get('carga_horaria_hr', '')} H/R
Série: {uc.get('serie', '')}
Objetivo Geral: {uc.get('objetivo_geral', '')}

CAPACIDADES:
{capacidades}

CONHECIMENTOS:
{conhecimentos}
"""

    prompts = {
        "plano_aulas": f"""Você é um especialista pedagógico do SENAI criando um plano de aulas detalhado.

{dados_uc}

{instrucao_formato}

Gere um plano de aulas com EXATAMENTE {num_aulas} aulas de 50 minutos (1 H/A cada).
Distribua os conhecimentos proporcionalmente ao longo das {num_aulas} aulas.
Cada aula deve ter: número, título, duração (50min), objetivos, conteúdo, 
metodologia SENAI (situação-problema/estudo de caso/projeto), recursos e avaliação.
Separe cada aula com ---
Siga a Metodologia SENAI de Educação Profissional com foco em competências.""",

        "apostila": f"""Você é um especialista técnico do SENAI criando uma apostila didática completa.

{dados_uc}

REFERÊNCIAS BIBLIOGRÁFICAS:
{referencias}

{instrucao_formato}

Crie uma apostila didática completa cobrindo TODOS os conhecimentos listados.
Inclua: apresentação, capítulos por tema, exemplos práticos, situações do mundo do trabalho.
Use linguagem clara e didática adequada ao nível técnico.
Finalize com as referências bibliográficas fornecidas.""",

        "atividades": f"""Você é um especialista pedagógico do SENAI criando atividades de fixação.

{dados_uc}

{instrucao_formato}

Crie um caderno de atividades com:
- 10 questões objetivas de múltipla escolha (4 alternativas cada)
- 5 questões dissertativas
- 2 situações-problema contextualizadas com o mercado de trabalho
Todas as questões devem avaliar as capacidades listadas.
Inclua gabarito das questões objetivas ao final.""",

        "avaliacao": f"""Você é um especialista pedagógico do SENAI criando uma avaliação formal.

{dados_uc}

{instrucao_formato}

Crie uma avaliação formal seguindo o padrão SENAI (100 pontos, 2 horas):
- Parte I: 10 questões objetivas (40 pts - 4 pts cada)
- Parte II: 3 questões dissertativas (30 pts - 10 pts cada)  
- Parte III: 1 situação-problema (30 pts)
Inclua campos de identificação do aluno, instruções e gabarito separado.
Questões devem avaliar as capacidades e conhecimentos da UC."""
    }

    return prompts.get(tipo, "")


# ─── Geradores principais ─────────────────────────────────────────────────────

def gerar_documento(uc_id: int, tipo: str, api_key: str, template_bytes: bytes = None) -> bytes:
    """
    Função unificada para gerar qualquer tipo de documento.
    
    uc_id: ID da UC no banco
    tipo: 'plano_aulas', 'apostila', 'atividades', 'avaliacao'
    api_key: chave da API Claude
    template_bytes: bytes do .docx template (opcional)
    """
    uc = carregar_uc(uc_id)
    if not uc:
        raise ValueError(f"UC {uc_id} não encontrada no banco")

    # Extrai texto do template se fornecido
    template_texto = ""
    if template_bytes:
        template_texto = extrair_texto_template(template_bytes)

    # Monta e envia o prompt
    client = get_client(api_key)
    prompt = montar_prompt(tipo, uc, template_texto)

    message = client.messages.create(
        model=MODEL,
        max_tokens=6000,
        messages=[{"role": "user", "content": prompt}]
    )
    conteudo_ia = message.content[0].text

    # Se tem template, tenta preservar a estrutura do Word original
    if template_bytes and template_texto:
        return gerar_com_template(conteudo_ia, template_bytes, uc, tipo)
    else:
        titulos = {
            "plano_aulas": "PLANO DE AULAS",
            "apostila": "APOSTILA",
            "atividades": "CADERNO DE ATIVIDADES",
            "avaliacao": "AVALIAÇÃO",
        }
        return texto_para_docx_padrao(
            conteudo_ia,
            titulos.get(tipo, tipo.upper()),
            f"Unidade Curricular: {uc['nome']}"
        )


def gerar_com_template(conteudo_ia: str, template_bytes: bytes, uc: dict, tipo: str) -> bytes:
    """
    Gera documento usando o template original como base estrutural.
    Cria novo doc baseado no template e insere o conteúdo gerado.
    """
    try:
        # Abre o template original como base
        doc = Document(io.BytesIO(template_bytes))

        # Limpa o conteúdo existente mantendo a estrutura
        for para in doc.paragraphs:
            if para.text.strip():
                para.clear()

        # Adiciona cabeçalho com nome da UC
        first_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        first_para.clear()
        run = first_para.add_run(f"{uc['nome']} — {tipo.replace('_', ' ').title()}")
        run.bold = True

        # Adiciona conteúdo gerado
        doc.add_paragraph()
        linhas = conteudo_ia.split("\n")
        for linha in linhas:
            if linha.strip():
                doc.add_paragraph(linha.strip())

        return doc_para_bytes(doc)
    except Exception:
        # Fallback: usa formato padrão se der erro com o template
        return texto_para_docx_padrao(conteudo_ia, tipo.upper(), uc['nome'])


# ─── Funções de compatibilidade (mantém chamadas antigas funcionando) ─────────

def gerar_plano_aulas(uc_id: int, api_key: str, template_bytes: bytes = None) -> bytes:
    return gerar_documento(uc_id, "plano_aulas", api_key, template_bytes)

def gerar_apostila(uc_id: int, api_key: str, template_bytes: bytes = None) -> bytes:
    return gerar_documento(uc_id, "apostila", api_key, template_bytes)

def gerar_atividades(uc_id: int, api_key: str, template_bytes: bytes = None) -> bytes:
    return gerar_documento(uc_id, "atividades", api_key, template_bytes)

def gerar_avaliacao(uc_id: int, api_key: str, template_bytes: bytes = None) -> bytes:
    return gerar_documento(uc_id, "avaliacao", api_key, template_bytes)